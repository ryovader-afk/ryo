"""出張旅費精算書のWord文書を生成するモジュール"""

import io
import datetime
from docx import Document
from docx.shared import Pt, Cm, Mm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


def _build_transport_note(td):
    """交通費のオプション情報から備考テキストを構築"""
    parts = []
    if td.get("airline"):
        parts.append(td["airline"])
    if td.get("flight_no"):
        parts.append(f"便名:{td['flight_no']}")
    if td.get("seat_class"):
        parts.append(td["seat_class"])
    if td.get("line"):
        parts.append(td["line"])
    if td.get("taxi_reason"):
        parts.append(f"理由:{td['taxi_reason']}")
    if td.get("rental_company"):
        parts.append(td["rental_company"])
    if td.get("rental_reason"):
        parts.append(f"理由:{td['rental_reason']}")
    if td.get("note"):
        parts.append(td["note"])
    return "／".join(parts)


def set_cell_shading(cell, color):
    """セルの背景色を設定"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_border(cell, **kwargs):
    """セルの罫線を設定"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}/>')
    for edge, val in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{val.get("val", "single")}" '
            f'w:sz="{val.get("sz", "4")}" w:space="0" '
            f'w:color="{val.get("color", "000000")}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)


def add_formatted_paragraph(doc, text, font_size=10.5, bold=False, alignment=None, space_after=None, font_name="游明朝"):
    """書式付き段落を追加"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.bold = bold
    if alignment:
        p.alignment = alignment
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p


def format_currency(amount):
    """金額を¥マーク付きカンマ区切りで書式化"""
    return f"¥{amount:,}"


def format_date_jp(date_str):
    """日付を和暦風に書式化 (例: 令和6年4月1日)"""
    if not date_str:
        return ""
    try:
        d = datetime.datetime.strptime(date_str, "%Y-%m-%d")
        year = d.year
        if year >= 2019:
            wareki_year = year - 2018
            era = "令和"
        elif year >= 1989:
            wareki_year = year - 1988
            era = "平成"
        else:
            wareki_year = year - 1925
            era = "昭和"
        return f"{era}{wareki_year}年{d.month}月{d.day}日"
    except ValueError:
        return date_str


def generate_travel_expense_report(
    company_name,
    employee_name,
    position,
    department,
    destination,
    purpose,
    purpose_detail,
    start_date,
    end_date,
    nights,
    accommodation_fee,
    daily_allowance,
    transport_details,
    trip_reason,
    trip_report,
):
    """出張旅費精算書を生成してバッファを返す"""
    doc = Document()

    # ページ設定 (A4)
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # デフォルトフォント設定
    style = doc.styles["Normal"]
    font = style.font
    font.name = "游明朝"
    font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "游明朝")

    days = nights + 1
    total_accommodation = accommodation_fee * nights
    total_allowance = daily_allowance * days
    transport_total = sum(t.get("amount", 0) for t in transport_details)
    grand_total = total_accommodation + total_allowance + transport_total

    start_date_jp = format_date_jp(start_date)
    end_date_jp = format_date_jp(end_date)
    today_jp = format_date_jp(datetime.date.today().strftime("%Y-%m-%d"))

    # === タイトル ===
    title = add_formatted_paragraph(doc, "出張旅費精算書", font_size=18, bold=True,
                                     alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)

    # === 作成日・会社名 ===
    add_formatted_paragraph(doc, f"作成日：{today_jp}", font_size=9,
                            alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_after=2)
    add_formatted_paragraph(doc, company_name, font_size=11, bold=True,
                            alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_after=6)

    # === 基本情報テーブル ===
    info_table = doc.add_table(rows=4, cols=4)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_table.style = "Table Grid"

    info_data = [
        ("所属部署", department, "役職", position),
        ("出張者氏名", employee_name, "印", ""),
        ("出張期間", f"{start_date_jp}　～　{end_date_jp}", "泊数", f"{nights}泊{days}日"),
        ("出張先", destination, "出張目的", purpose),
    ]

    for i, (h1, v1, h2, v2) in enumerate(info_data):
        row = info_table.rows[i]
        for j, (header, value) in enumerate([(h1, v1), (h2, v2)]):
            hcell = row.cells[j * 2]
            vcell = row.cells[j * 2 + 1]

            hcell.text = ""
            hp = hcell.paragraphs[0]
            hr = hp.add_run(header)
            hr.font.size = Pt(9)
            hr.font.name = "游ゴシック"
            hr._element.rPr.rFonts.set(qn("w:eastAsia"), "游ゴシック")
            hr.bold = True
            hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_shading(hcell, "E8F0FE")

            vcell.text = ""
            vp = vcell.paragraphs[0]
            vr = vp.add_run(value)
            vr.font.size = Pt(10)
            vr.font.name = "游明朝"
            vr._element.rPr.rFonts.set(qn("w:eastAsia"), "游明朝")

    # 列幅設定
    for row in info_table.rows:
        row.cells[0].width = Cm(2.8)
        row.cells[1].width = Cm(5.5)
        row.cells[2].width = Cm(2.8)
        row.cells[3].width = Cm(4.0)

    doc.add_paragraph()

    # === 出張理由 ===
    add_formatted_paragraph(doc, "■ 出張理由", font_size=11, bold=True, space_after=2)
    reason_table = doc.add_table(rows=1, cols=1)
    reason_table.style = "Table Grid"
    reason_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = reason_table.rows[0].cells[0]
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(trip_reason)
    r.font.size = Pt(10)
    r.font.name = "游明朝"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "游明朝")

    if purpose_detail:
        p2 = cell.add_paragraph()
        r2 = p2.add_run(f"\n【詳細】{purpose_detail}")
        r2.font.size = Pt(10)
        r2.font.name = "游明朝"
        r2._element.rPr.rFonts.set(qn("w:eastAsia"), "游明朝")

    doc.add_paragraph()

    # === 旅費明細テーブル ===
    add_formatted_paragraph(doc, "■ 旅費明細", font_size=11, bold=True, space_after=2)

    # 交通費テーブル
    if transport_details:
        add_formatted_paragraph(doc, "【交通費】", font_size=10, bold=True, space_after=1)

        # 備考がある交通費があるかチェック
        has_notes = any(_build_transport_note(td) for td in transport_details)
        col_count = 5 if has_notes else 4

        t_table = doc.add_table(rows=1 + len(transport_details), cols=col_count)
        t_table.style = "Table Grid"
        t_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        headers = ["日付", "区間", "交通機関", "金額"]
        if has_notes:
            headers.append("備考")
        for j, h in enumerate(headers):
            cell = t_table.rows[0].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            r = p.add_run(h)
            r.font.size = Pt(9)
            r.font.name = "游ゴシック"
            r._element.rPr.rFonts.set(qn("w:eastAsia"), "游ゴシック")
            r.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_shading(cell, "E8F0FE")

        for i, td in enumerate(transport_details):
            row = t_table.rows[i + 1]
            values = [
                format_date_jp(td.get("date", "")),
                td.get("route", ""),
                td.get("method", ""),
                format_currency(td.get("amount", 0)),
            ]
            if has_notes:
                values.append(_build_transport_note(td))
            for j, val in enumerate(values):
                cell = row.cells[j]
                cell.text = ""
                p = cell.paragraphs[0]
                r = p.add_run(val)
                r.font.size = Pt(10 if j != 4 else 9)
                r.font.name = "游明朝"
                r._element.rPr.rFonts.set(qn("w:eastAsia"), "游明朝")
                if j == 3:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        doc.add_paragraph()

    # 宿泊費・日当テーブル
    add_formatted_paragraph(doc, "【宿泊費・日当】", font_size=10, bold=True, space_after=1)
    s_table = doc.add_table(rows=4, cols=4)
    s_table.style = "Table Grid"
    s_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    summary_headers = ["項目", "単価", "数量", "金額"]
    for j, h in enumerate(summary_headers):
        cell = s_table.rows[0].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.font.size = Pt(9)
        r.font.name = "游ゴシック"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "游ゴシック")
        r.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, "E8F0FE")

    summary_data = [
        ("宿泊費", format_currency(accommodation_fee), f"{nights}泊", format_currency(total_accommodation)),
        ("日当", format_currency(daily_allowance), f"{days}日", format_currency(total_allowance)),
        ("交通費合計", "", "", format_currency(transport_total)),
    ]

    for i, (item, unit, qty, total) in enumerate(summary_data):
        row = s_table.rows[i + 1]
        values = [item, unit, qty, total]
        for j, val in enumerate(values):
            cell = row.cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            r = p.add_run(val)
            r.font.size = Pt(10)
            r.font.name = "游明朝"
            r._element.rPr.rFonts.set(qn("w:eastAsia"), "游明朝")
            if j == 3:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_paragraph()

    # === 合計金額 ===
    total_table = doc.add_table(rows=1, cols=2)
    total_table.style = "Table Grid"
    total_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    label_cell = total_table.rows[0].cells[0]
    label_cell.text = ""
    lp = label_cell.paragraphs[0]
    lr = lp.add_run("合計支給額")
    lr.font.size = Pt(12)
    lr.font.name = "游ゴシック"
    lr._element.rPr.rFonts.set(qn("w:eastAsia"), "游ゴシック")
    lr.bold = True
    lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_shading(label_cell, "1F4E79")
    lr.font.color.rgb = RGBColor(255, 255, 255)

    amount_cell = total_table.rows[0].cells[1]
    amount_cell.text = ""
    ap = amount_cell.paragraphs[0]
    ar = ap.add_run(format_currency(grand_total))
    ar.font.size = Pt(14)
    ar.font.name = "游ゴシック"
    ar._element.rPr.rFonts.set(qn("w:eastAsia"), "游ゴシック")
    ar.bold = True
    ap.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    label_cell.width = Cm(5)
    amount_cell.width = Cm(10)

    doc.add_paragraph()

    # === 出張報告 ===
    add_formatted_paragraph(doc, "■ 出張報告", font_size=11, bold=True, space_after=2)
    report_table = doc.add_table(rows=1, cols=1)
    report_table.style = "Table Grid"
    report_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    rcell = report_table.rows[0].cells[0]
    rcell.text = ""
    rp = rcell.paragraphs[0]
    rr = rp.add_run(trip_report)
    rr.font.size = Pt(10)
    rr.font.name = "游明朝"
    rr._element.rPr.rFonts.set(qn("w:eastAsia"), "游明朝")

    doc.add_paragraph()

    # === 承認欄 ===
    add_formatted_paragraph(doc, "■ 承認", font_size=11, bold=True, space_after=2)
    a_table = doc.add_table(rows=2, cols=4)
    a_table.style = "Table Grid"
    a_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    approval_headers = ["申請者", "所属長", "経理部", "代表取締役"]
    for j, h in enumerate(approval_headers):
        cell = a_table.rows[0].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.font.size = Pt(9)
        r.font.name = "游ゴシック"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "游ゴシック")
        r.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, "E8F0FE")

    for j in range(4):
        cell = a_table.rows[1].cells[j]
        cell.text = ""
        # 印鑑スペースの高さ確保
        p = cell.paragraphs[0]
        for _ in range(3):
            p.add_run("\n")

    doc.add_paragraph()

    # === 備考 ===
    add_formatted_paragraph(doc, "備考：本精算書は出張旅費規程に基づき作成されています。",
                            font_size=8, space_after=1)
    add_formatted_paragraph(doc, "※宿泊費は規程に定める限度額の範囲内で実費精算とします。",
                            font_size=8, space_after=1)
    add_formatted_paragraph(doc, "※交通費は領収書を添付してください。",
                            font_size=8, space_after=0)

    # バッファに書き出し
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
