#!/usr/bin/env python3
"""不動産投資キャッシュフローシミュレーション結果のエクスポート（PDF/Excel/Word）"""

import io
import datetime
from typing import Any, Dict, Union


def _fmt_yen(n: Union[int, float]) -> str:
    if n < 0:
        return f"-¥{abs(int(n)):,}"
    return f"¥{int(n):,}"


def _fmt_pct(n: float) -> str:
    return f"{n:.2f}%"


# ──────────────────────────────────────────────
# Excel
# ──────────────────────────────────────────────
def generate_excel(data: Dict[str, Any]) -> io.BytesIO:
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, PatternFill, Border, Side, numbers

    wb = Workbook()

    header_font = Font(name="Yu Gothic", bold=True, size=12)
    title_font = Font(name="Yu Gothic", bold=True, size=14)
    normal_font = Font(name="Yu Gothic", size=11)
    yen_fmt = '#,##0;-#,##0'
    pct_fmt = '0.00"%"'

    header_fill = PatternFill(start_color="1E3A5F", end_color="1E3A5F", fill_type="solid")
    header_font_w = Font(name="Yu Gothic", bold=True, size=11, color="FFFFFF")
    positive_font = Font(name="Yu Gothic", size=11, color="008000")
    negative_font = Font(name="Yu Gothic", size=11, color="CC0000")
    thin_border = Border(
        left=Side(style="thin"), right=Side(style="thin"),
        top=Side(style="thin"), bottom=Side(style="thin"),
    )

    summary = data["summary"]
    monthly = data["monthly"]
    annual = data["annual"]
    prop = data["property"]

    # --- Sheet 1: サマリー ---
    ws = wb.active
    ws.title = "サマリー"
    ws.column_dimensions["A"].width = 30
    ws.column_dimensions["B"].width = 25

    ws.merge_cells("A1:B1")
    ws["A1"] = f"不動産投資キャッシュフロー - {prop.get('name', '')}"
    ws["A1"].font = title_font

    ws["A2"] = f"作成日: {datetime.date.today().isoformat()}"
    ws["A2"].font = Font(name="Yu Gothic", size=10, color="888888")

    rows = [
        ("物件情報", ""),
        ("物件価格", prop.get("totalPrice", 0)),
        ("建物価格（税抜）", prop.get("buildingPrice", 0)),
        ("築年数", f"{prop.get('buildingAge', 0)}年"),
        ("建物完成月", f"{prop.get('completionMonth', 1)}月目"),
        ("借入金額", prop.get("loanAmount", 0)),
        ("金利", f"{prop.get('interestRate', 0)}%"),
        ("借入期間", f"{prop.get('loanTermYears', 0)}年"),
        ("据え置き期間", f"{prop.get('gracePeriod', 0)}ヶ月"),
        ("", ""),
        ("収支指標", ""),
        ("表面利回り", summary.get("grossYield", "")),
        ("実質利回り", summary.get("netYield", "")),
        ("月額返済額", summary.get("monthlyPayment", 0)),
        ("初年度キャッシュフロー", summary.get("year1CF", 0)),
        ("消費税還付額", summary.get("refund", 0)),
        ("還付込み初年度CF", summary.get("year1CFWithRefund", 0)),
        ("据え置き期間利益効果", summary.get("graceBenefit", 0)),
        ("投資回収年数", summary.get("breakeven", "")),
    ]

    for i, (label, value) in enumerate(rows, start=4):
        ws.cell(row=i, column=1, value=label).font = header_font if value == "" and label else normal_font
        cell = ws.cell(row=i, column=2)
        if isinstance(value, (int, float)) and label:
            cell.value = value
            cell.number_format = yen_fmt
        else:
            cell.value = value
        cell.font = normal_font

    # --- Sheet 2: 月別CF ---
    ws2 = wb.create_sheet("月別キャッシュフロー")
    headers = ["月", "家賃収入", "その他収入", "ローン返済", "経費", "月間CF", "累計CF"]
    col_widths = [8, 16, 16, 16, 16, 16, 16]

    for j, (h, w) in enumerate(zip(headers, col_widths), start=1):
        cell = ws2.cell(row=1, column=j, value=h)
        cell.font = header_font_w
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center")
        cell.border = thin_border
        ws2.column_dimensions[cell.column_letter].width = w

    for i, row in enumerate(monthly, start=2):
        values = [
            f"{row['month']}月",
            row["rent"], row["other"], row["loan"], row["expense"], row["cf"], row["cumCF"],
        ]
        for j, v in enumerate(values, start=1):
            cell = ws2.cell(row=i, column=j, value=v)
            cell.border = thin_border
            if j >= 2:
                cell.number_format = yen_fmt
                if j in (4, 5):
                    cell.font = negative_font
                elif j == 6:
                    cell.font = positive_font if v >= 0 else negative_font
                elif j == 7:
                    cell.font = positive_font if v >= 0 else negative_font
                else:
                    cell.font = normal_font
            else:
                cell.font = normal_font
                cell.alignment = Alignment(horizontal="center")

    # --- Sheet 3: 年別CF ---
    ws3 = wb.create_sheet("年別キャッシュフロー")
    headers3 = ["年", "年間収入", "年間ローン返済", "年間経費", "年間CF", "累計CF"]
    col_widths3 = [10, 18, 18, 18, 18, 18]

    for j, (h, w) in enumerate(zip(headers3, col_widths3), start=1):
        cell = ws3.cell(row=1, column=j, value=h)
        cell.font = header_font_w
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center")
        cell.border = thin_border
        ws3.column_dimensions[cell.column_letter].width = w

    for i, row in enumerate(annual, start=2):
        values = [
            f"{row['year']}年目",
            row["income"], row["loan"], row["expense"], row["cf"], row["cumCF"],
        ]
        for j, v in enumerate(values, start=1):
            cell = ws3.cell(row=i, column=j, value=v)
            cell.border = thin_border
            if j >= 2:
                cell.number_format = yen_fmt
                if j in (3, 4):
                    cell.font = negative_font
                elif j >= 5:
                    cell.font = positive_font if v >= 0 else negative_font
                else:
                    cell.font = normal_font
            else:
                cell.font = normal_font
                cell.alignment = Alignment(horizontal="center")

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf


# ──────────────────────────────────────────────
# Word
# ──────────────────────────────────────────────
def generate_word(data: Dict[str, Any]) -> io.BytesIO:
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Yu Gothic"
    style.font.size = Pt(10.5)

    summary = data["summary"]
    monthly = data["monthly"]
    annual = data["annual"]
    prop = data["property"]

    # Title
    title = doc.add_heading(level=1)
    run = title.add_run(f"不動産投資キャッシュフロー分析レポート")
    run.font.size = Pt(18)

    doc.add_paragraph(f"物件名: {prop.get('name', '未設定')}")
    doc.add_paragraph(f"作成日: {datetime.date.today().isoformat()}")

    # Property info
    doc.add_heading("物件概要", level=2)
    info_data = [
        ("物件価格", _fmt_yen(prop.get("totalPrice", 0))),
        ("建物価格（税抜）", _fmt_yen(prop.get("buildingPrice", 0))),
        ("築年数", f"{prop.get('buildingAge', 0)}年"),
        ("建物完成月", f"{prop.get('completionMonth', 1)}月目"),
        ("借入金額", _fmt_yen(prop.get("loanAmount", 0))),
        ("金利", f"{prop.get('interestRate', 0)}%"),
        ("借入期間", f"{prop.get('loanTermYears', 0)}年"),
        ("据え置き期間", f"{prop.get('gracePeriod', 0)}ヶ月（返済¥0）"),
        ("月額返済額", _fmt_yen(summary.get("monthlyPayment", 0))),
    ]
    info_table = doc.add_table(rows=len(info_data), cols=2, style="Light Grid Accent 1")
    for i, (label, value) in enumerate(info_data):
        info_table.cell(i, 0).text = label
        info_table.cell(i, 1).text = str(value)

    # Summary
    doc.add_heading("収支サマリー", level=2)
    sum_table = doc.add_table(rows=7, cols=2, style="Light Grid Accent 1")
    sum_data = [
        ("表面利回り", summary.get("grossYield", "-")),
        ("実質利回り", summary.get("netYield", "-")),
        ("初年度キャッシュフロー", _fmt_yen(summary.get("year1CF", 0))),
        ("消費税還付額", _fmt_yen(summary.get("refund", 0))),
        ("還付込み初年度CF", _fmt_yen(summary.get("year1CFWithRefund", 0))),
        ("据え置き利益効果", _fmt_yen(summary.get("graceBenefit", 0))),
        ("投資回収年数", summary.get("breakeven", "-")),
    ]
    for i, (label, value) in enumerate(sum_data):
        sum_table.cell(i, 0).text = label
        sum_table.cell(i, 1).text = str(value)

    # Monthly CF
    doc.add_heading("初年度 月別キャッシュフロー", level=2)
    cols = ["月", "家賃収入", "その他収入", "ローン返済", "経費", "月間CF", "累計CF"]
    m_table = doc.add_table(rows=len(monthly) + 1, cols=7, style="Light Grid Accent 1")
    m_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for j, h in enumerate(cols):
        cell = m_table.cell(0, j)
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True

    for i, row in enumerate(monthly, start=1):
        m_table.cell(i, 0).text = f"{row['month']}月"
        m_table.cell(i, 1).text = _fmt_yen(row["rent"])
        m_table.cell(i, 2).text = _fmt_yen(row["other"])
        m_table.cell(i, 3).text = _fmt_yen(row["loan"])
        m_table.cell(i, 4).text = _fmt_yen(row["expense"])
        m_table.cell(i, 5).text = _fmt_yen(row["cf"])
        m_table.cell(i, 6).text = _fmt_yen(row["cumCF"])

    # Annual CF
    doc.add_heading("年別キャッシュフロー", level=2)
    cols2 = ["年", "年間収入", "年間ローン返済", "年間経費", "年間CF", "累計CF"]
    a_table = doc.add_table(rows=len(annual) + 1, cols=6, style="Light Grid Accent 1")
    a_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for j, h in enumerate(cols2):
        cell = a_table.cell(0, j)
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True

    for i, row in enumerate(annual, start=1):
        a_table.cell(i, 0).text = f"{row['year']}年目"
        a_table.cell(i, 1).text = _fmt_yen(row["income"])
        a_table.cell(i, 2).text = _fmt_yen(row["loan"])
        a_table.cell(i, 3).text = _fmt_yen(row["expense"])
        a_table.cell(i, 4).text = _fmt_yen(row["cf"])
        a_table.cell(i, 5).text = _fmt_yen(row["cumCF"])

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ──────────────────────────────────────────────
# PDF
# ──────────────────────────────────────────────
def generate_pdf(data: Dict[str, Any]) -> io.BytesIO:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.platypus import (
        SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak,
    )
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    import os

    # Try to register a Japanese font
    jp_font = "Helvetica"
    jp_font_bold = "Helvetica-Bold"
    font_paths = [
        ("/System/Library/Fonts/ヒラギノ角ゴシック W3.ttc", "HiraginoW3"),
        ("/System/Library/Fonts/ヒラギノ角ゴシック W6.ttc", "HiraginoW6"),
        ("/System/Library/Fonts/Hiragino Sans GB.ttc", "HiraginoSans"),
        ("/Library/Fonts/Arial Unicode.ttf", "ArialUnicode"),
    ]
    for path, name in font_paths:
        if os.path.exists(path):
            try:
                pdfmetrics.registerFont(TTFont(name, path, subfontIndex=0))
                jp_font = name
                jp_font_bold = name
                break
            except Exception:
                continue

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, topMargin=20*mm, bottomMargin=15*mm,
                            leftMargin=15*mm, rightMargin=15*mm)

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="JP", fontName=jp_font, fontSize=10, leading=14))
    styles.add(ParagraphStyle(name="JPTitle", fontName=jp_font_bold, fontSize=18, leading=24,
                               spaceAfter=6))
    styles.add(ParagraphStyle(name="JPHeading", fontName=jp_font_bold, fontSize=13, leading=18,
                               spaceAfter=4, spaceBefore=12,
                               textColor=colors.HexColor("#1E3A5F")))
    styles.add(ParagraphStyle(name="JPSmall", fontName=jp_font, fontSize=8, leading=10,
                               textColor=colors.gray))

    summary = data["summary"]
    monthly = data["monthly"]
    annual = data["annual"]
    prop = data["property"]

    elements = []

    # Title
    elements.append(Paragraph("不動産投資キャッシュフロー分析レポート", styles["JPTitle"]))
    elements.append(Paragraph(f"物件名: {prop.get('name', '未設定')}　　作成日: {datetime.date.today().isoformat()}", styles["JPSmall"]))
    elements.append(Spacer(1, 8*mm))

    # Summary table
    elements.append(Paragraph("収支サマリー", styles["JPHeading"]))

    dark_blue = colors.HexColor("#1E3A5F")
    light_bg = colors.HexColor("#F0F4F8")

    sum_data = [
        ["指標", "値"],
        ["表面利回り", summary.get("grossYield", "-")],
        ["実質利回り", summary.get("netYield", "-")],
        ["月額返済額", _fmt_yen(summary.get("monthlyPayment", 0))],
        ["初年度CF", _fmt_yen(summary.get("year1CF", 0))],
        ["消費税還付額", _fmt_yen(summary.get("refund", 0))],
        ["還付込み初年度CF", _fmt_yen(summary.get("year1CFWithRefund", 0))],
        ["据え置き利益効果", _fmt_yen(summary.get("graceBenefit", 0))],
        ["投資回収年数", summary.get("breakeven", "-")],
    ]
    sum_table = Table(sum_data, colWidths=[60*mm, 50*mm])
    sum_table.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (-1, -1), jp_font),
        ("FONTSIZE", (0, 0), (-1, -1), 10),
        ("BACKGROUND", (0, 0), (-1, 0), dark_blue),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), jp_font_bold),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, light_bg]),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#CCCCCC")),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    elements.append(sum_table)
    elements.append(Spacer(1, 8*mm))

    # Monthly table
    elements.append(Paragraph("初年度 月別キャッシュフロー", styles["JPHeading"]))
    m_headers = ["月", "家賃収入", "その他", "ローン返済", "経費", "月間CF", "累計CF"]
    m_data = [m_headers]
    for row in monthly:
        m_data.append([
            f"{row['month']}月",
            _fmt_yen(row["rent"]),
            _fmt_yen(row["other"]),
            _fmt_yen(row["loan"]),
            _fmt_yen(row["expense"]),
            _fmt_yen(row["cf"]),
            _fmt_yen(row["cumCF"]),
        ])
    m_table = Table(m_data, colWidths=[14*mm, 26*mm, 22*mm, 26*mm, 22*mm, 26*mm, 26*mm])
    m_table.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (-1, -1), jp_font),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("BACKGROUND", (0, 0), (-1, 0), dark_blue),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), jp_font_bold),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, light_bg]),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#CCCCCC")),
        ("ALIGN", (1, 0), (-1, -1), "RIGHT"),
        ("ALIGN", (0, 0), (0, -1), "CENTER"),
        ("TOPPADDING", (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
    ]))
    elements.append(m_table)

    # Page break before annual
    elements.append(PageBreak())

    # Annual table
    elements.append(Paragraph("年別キャッシュフロー", styles["JPHeading"]))
    a_headers = ["年", "年間収入", "年間ローン返済", "年間経費", "年間CF", "累計CF"]
    a_data = [a_headers]
    for row in annual:
        a_data.append([
            f"{row['year']}年目",
            _fmt_yen(row["income"]),
            _fmt_yen(row["loan"]),
            _fmt_yen(row["expense"]),
            _fmt_yen(row["cf"]),
            _fmt_yen(row["cumCF"]),
        ])
    a_table = Table(a_data, colWidths=[16*mm, 30*mm, 30*mm, 28*mm, 28*mm, 30*mm])
    a_table.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (-1, -1), jp_font),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("BACKGROUND", (0, 0), (-1, 0), dark_blue),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), jp_font_bold),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, light_bg]),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#CCCCCC")),
        ("ALIGN", (1, 0), (-1, -1), "RIGHT"),
        ("ALIGN", (0, 0), (0, -1), "CENTER"),
        ("TOPPADDING", (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
    ]))
    elements.append(a_table)

    doc.build(elements)
    buf.seek(0)
    return buf
